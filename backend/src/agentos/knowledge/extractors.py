"""Format-specific extractors for local Knowledge Vault documents."""

import mimetypes
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any


class UnsupportedDocumentError(ValueError):
    """Raised when no safe extractor exists for a document format."""


@dataclass(frozen=True)
class ExtractedBlock:
    """Normalized text and source location emitted by an extractor."""

    text: str
    heading_path: list[str]
    page_number: int | None = None
    sheet_name: str | None = None
    source_location: str | None = None
    block_type: str = "paragraph"


@dataclass(frozen=True)
class ExtractedDocument:
    """The normalized result of extracting one source document."""

    path: Path
    mime_type: str
    blocks: list[ExtractedBlock]
    structure: dict[str, Any]


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
_MIME_TYPES = {
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".txt": "text/plain",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def extract_document(path: Path) -> ExtractedDocument:
    """Extract supported local document formats into citation-aware blocks."""
    path = Path(path)
    suffix = path.suffix.lower()
    mime_type = _MIME_TYPES.get(suffix) or mimetypes.guess_type(path.name)[0]
    if mime_type not in set(_MIME_TYPES.values()):
        raise UnsupportedDocumentError(
            f"No extractor for {mime_type or suffix}; image OCR is not supported yet"
        )

    if suffix in {".md", ".markdown", ".txt"}:
        blocks = _extract_text(path)
    elif suffix == ".pdf":
        blocks = _extract_pdf(path)
    elif suffix == ".docx":
        blocks = _extract_docx(path)
    else:
        blocks = _extract_xlsx(path)
    return ExtractedDocument(
        path=path,
        mime_type=mime_type,
        blocks=blocks,
        structure=_build_structure(blocks),
    )


def _build_structure(blocks: list[ExtractedBlock]) -> dict[str, Any]:
    """Build compact structural metadata without copying document contents."""
    sections: list[list[str]] = []
    pages: set[int] = set()
    tables: list[str] = []
    images: list[str] = []
    sheets: set[str] = set()
    for block in blocks:
        if block.heading_path and block.heading_path not in sections:
            sections.append(block.heading_path.copy())
        if block.page_number is not None:
            pages.add(block.page_number)
        if block.sheet_name:
            sheets.add(block.sheet_name)
        if block.source_location:
            if block.source_location.startswith("table "):
                tables.append(block.source_location)
            elif block.source_location.startswith("image "):
                images.append(block.source_location)
    return {
        "sections": sections,
        "pages": sorted(pages),
        "tables": tables,
        "images": images,
        "sheets": sorted(sheets),
    }


def _extract_text(path: Path) -> list[ExtractedBlock]:
    heading_stack: list[str] = []
    blocks: list[ExtractedBlock] = []
    paragraph: list[str] = []

    def flush() -> None:
        if paragraph:
            blocks.append(ExtractedBlock(" ".join(paragraph), heading_stack.copy()))
            paragraph.clear()

    for line in path.read_text(encoding="utf-8").splitlines():
        match = _HEADING_RE.match(line) if path.suffix.lower() != ".txt" else None
        if match:
            flush()
            heading_stack[:] = heading_stack[: len(match.group(1)) - 1]
            heading_stack.append(match.group(2))
        elif line.strip():
            paragraph.append(line.strip())
        else:
            flush()
    flush()
    return blocks


def _extract_pdf(path: Path) -> list[ExtractedBlock]:
    from pypdf import PdfReader

    reader = PdfReader(path)
    outline_paths = _pdf_outline_paths(reader)
    blocks: list[ExtractedBlock] = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        blocks.append(
            ExtractedBlock(
                text=text,
                heading_path=outline_paths.get(page_number, []),
                page_number=page_number,
                block_type="page",
            )
        )
    return blocks


def _pdf_outline_paths(reader: Any) -> dict[int, list[str]]:
    """Map each PDF page to the deepest bookmark path that starts on it."""
    entries: list[tuple[int, list[str]]] = []

    def visit(items: list[Any], parents: list[str]) -> None:
        current_parents = parents
        for item in items:
            if isinstance(item, list):
                visit(item, current_parents)
                continue
            try:
                page_number = reader.get_destination_page_number(item) + 1
            except (TypeError, ValueError, IndexError):
                continue
            title = str(getattr(item, "title", "")).strip()
            if not title:
                continue
            current_parents = [*parents, title]
            entries.append((page_number, current_parents))

    visit(reader.outline or [], [])
    entries.sort(key=lambda entry: (entry[0], len(entry[1])))
    return {
        page_number: max(
            (entry for entry in entries if entry[0] <= page_number),
            key=lambda entry: (entry[0], len(entry[1])),
            default=(0, []),
        )[1]
        for page_number in range(1, len(reader.pages) + 1)
    }


def _extract_docx(path: Path) -> list[ExtractedBlock]:
    from docx import Document as WordDocument

    source = WordDocument(path)
    heading_stack: list[str] = []
    blocks: list[ExtractedBlock] = []
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name or ""
        if style_name.startswith("Heading"):
            level = int(style_name.removeprefix("Heading").strip() or "1")
            heading_stack[:] = heading_stack[: level - 1]
            heading_stack.append(text)
        else:
            blocks.append(ExtractedBlock(text, heading_stack.copy()))

    for table_index, table in enumerate(source.tables, start=1):
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if rows:
            blocks.append(
                ExtractedBlock(
                    text="\n".join(rows),
                    heading_path=heading_stack.copy(),
                    source_location=f"table {table_index}",
                    block_type="table",
                )
            )

    with zipfile.ZipFile(path) as archive:
        images = sorted(name for name in archive.namelist() if name.startswith("word/media/"))
    for image_index, image_name in enumerate(images, start=1):
        blocks.append(
            ExtractedBlock(
                text=f"Embedded image: {Path(image_name).name}",
                heading_path=heading_stack.copy(),
                source_location=f"image {image_index}",
                block_type="figure",
            )
        )
    return blocks


def _extract_xlsx(path: Path) -> list[ExtractedBlock]:
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter

    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[ExtractedBlock] = []
    for sheet in workbook.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        values = [" | ".join("" if value is None else str(value) for value in row) for row in rows]
        last_row = len(rows)
        last_column = max(len(row) for row in rows)
        location = f"{sheet.title}!A1:{get_column_letter(last_column)}{last_row}"
        blocks.append(
            ExtractedBlock(
                text="\n".join(values),
                heading_path=[],
                sheet_name=sheet.title,
                source_location=location,
                block_type="sheet",
            )
        )
    return blocks
