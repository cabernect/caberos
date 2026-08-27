from pathlib import Path

import pytest

from agentos.knowledge.chunker import chunk_extracted_blocks
from agentos.knowledge.extractors import (
    ExtractedBlock,
    UnsupportedDocumentError,
    extract_document,
)


def test_markdown_extraction_preserves_heading_context(tmp_path: Path):
    path = tmp_path / "guide.md"
    path.write_text("# Setup\n\nInstall CaberOS.\n\n## Config\n\nSet the key.", encoding="utf-8")

    document = extract_document(path)

    assert [block.heading_path for block in document.blocks] == [
        ["Setup"],
        ["Setup", "Config"],
    ]
    assert document.blocks[0].text == "Install CaberOS."
    assert document.structure == {
        "sections": [["Setup"], ["Setup", "Config"]],
        "pages": [],
        "tables": [],
        "images": [],
        "sheets": [],
    }


def test_pdf_extraction_preserves_page_numbers(tmp_path: Path):
    pytest.importorskip("pypdf")
    from pypdf import PdfWriter

    path = tmp_path / "report.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    with path.open("wb") as output:
        writer.write(output)

    document = extract_document(path)

    assert document.mime_type == "application/pdf"
    assert document.blocks[0].page_number == 1
    assert document.blocks[0].block_type == "page"


def test_docx_extraction_preserves_heading_context(tmp_path: Path):
    pytest.importorskip("docx")
    from docx import Document as WordDocument

    path = tmp_path / "guide.docx"
    source = WordDocument()
    source.add_heading("Setup", level=1)
    source.add_paragraph("Install CaberOS.")
    source.add_heading("Config", level=2)
    source.add_paragraph("Set the key.")
    source.save(path)

    document = extract_document(path)

    assert [block.heading_path for block in document.blocks] == [
        ["Setup"],
        ["Setup", "Config"],
    ]


def test_xlsx_extraction_preserves_sheet_and_cell_ranges(tmp_path: Path):
    pytest.importorskip("openpyxl")
    from openpyxl import Workbook

    path = tmp_path / "budget.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Q3"
    sheet.append(["Item", "Amount"])
    sheet.append(["Hosting", 42])
    workbook.save(path)

    document = extract_document(path)

    assert document.blocks[0].sheet_name == "Q3"
    assert "Hosting" in document.blocks[0].text
    assert document.blocks[0].source_location == "Q3!A1:B2"


def test_chunking_preserves_pdf_source_metadata():
    blocks = [ExtractedBlock("Quarterly revenue", [], page_number=4, source_location="page 4")]

    chunks = chunk_extracted_blocks(blocks, max_tokens=10, overlap_tokens=0)

    assert chunks[0].page_number == 4
    assert chunks[0].source_location == "page 4"


def test_images_are_visible_as_unsupported_until_ocr_is_added(tmp_path: Path):
    path = tmp_path / "diagram.png"
    path.write_bytes(b"not an image")

    with pytest.raises(UnsupportedDocumentError, match="OCR"):
        extract_document(path)
